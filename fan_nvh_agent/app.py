import os
import io
import base64
import pandas as pd
import numpy as np
import requests
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32MB
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# ──────────────────────────────────────────────
# 工具函数
# ──────────────────────────────────────────────

def detect_anomalies_zscore(values, threshold=2.0):
    """Z-Score 异常检测：返回每个值的异常标记列表"""
    arr = np.array(values, dtype=float)
    mean = np.mean(arr)
    std = np.std(arr)
    if std == 0:
        return [False] * len(values)
    z_scores = np.abs((arr - mean) / std)
    return (z_scores > threshold).tolist()

def detect_anomalies_iqr(values, multiplier=1.5):
    """IQR 箱线图异常检测"""
    arr = np.array(values, dtype=float)
    q1 = np.percentile(arr, 25)
    q3 = np.percentile(arr, 75)
    iqr = q3 - q1
    lower = q1 - multiplier * iqr
    upper = q3 + multiplier * iqr
    return ((arr < lower) | (arr > upper)).tolist()

def compute_pareto_front(airflow, noise):
    """
    计算帕累托前沿（非支配排序）。
    风量越大越好（最大化），噪音越小越好（最小化）。
    返回每个点是否在 Pareto 前沿上的布尔列表。
    """
    n = len(airflow)
    pareto = [True] * n
    for i in range(n):
        for j in range(n):
            if i == j:
                continue
            # j 支配 i：j 风量 >= i 风量 且 j 噪音 <= i 噪音，且至少一个严格优于
            if airflow[j] >= airflow[i] and noise[j] <= noise[i]:
                if airflow[j] > airflow[i] or noise[j] < noise[i]:
                    pareto[i] = False
                    break
    return pareto

def fit_linear_model(df):
    """拟合噪音 ~ 转速 + 风量 的线性回归模型，返回系数"""
    from numpy.linalg import lstsq
    X = df[['转速(RPM)', '风量(CFM)']].values.astype(float)
    y = df['噪音(dB)'].values.astype(float)
    # 添加截距项
    X_with_intercept = np.column_stack([X, np.ones(len(X))])
    coeffs, residuals, rank, singular = lstsq(X_with_intercept, y, rcond=None)
    return {'rpm_coef': coeffs[0], 'airflow_coef': coeffs[1], 'intercept': coeffs[2]}

def predict_noise(rpm, airflow, model):
    """使用线性模型预测噪音"""
    return model['rpm_coef'] * rpm + model['airflow_coef'] * airflow + model['intercept']

# ──────────────────────────────────────────────
# 路由
# ──────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/process', methods=['POST'])
def process_data():
    """主分析接口：上传 CSV → 数据清洗 → 异常检测 → 图表数据 + AI 分析"""
    if 'file' not in request.files:
        return jsonify({'error': '未检测到文件上传'})

    file = request.files['file']
    search_keyword = request.form.get('keyword', '').strip()
    api_key = request.form.get('api_key', '').strip()
    api_url = request.form.get('api_url',
        'https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions').strip()
    api_model = request.form.get('api_model', 'qwen3.7-plus').strip()

    # 异常检测阈值参数
    zscore_threshold = float(request.form.get('zscore_threshold', 2.0))
    noise_alert_threshold = float(request.form.get('noise_alert', 0))  # 0 表示不启用
    airflow_alert_threshold = float(request.form.get('airflow_alert', 0))

    if file.filename == '':
        return jsonify({'error': '未选择文件'})

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filepath)

    try:
        df = pd.read_csv(filepath, encoding='gbk')

        # 清洗
        if '车型' in df.columns:
            df = df[~df['车型'].astype(str).str.contains('总计', na=False)]
        if '市占率' in df.columns:
            df = df.drop(columns=['市占率'], errors='ignore')

        if search_keyword and '车型' in df.columns:
            df = df[df['车型'].astype(str).str.contains(search_keyword, na=False)]

        if df.empty:
            return jsonify({'error': '没有符合条件的数据'})

        # ---------- 提取核心列 ----------
        fan_labels = df['风扇型号'].tolist() if '风扇型号' in df.columns else [f"样本{i}" for i in range(len(df))]
        airflow = [float(v) for v in df['风量(CFM)'].tolist()] if '风量(CFM)' in df.columns else []
        noise = [float(v) for v in df['噪音(dB)'].tolist()] if '噪音(dB)' in df.columns else []
        pass_rate = [float(v) for v in df['达标率'].tolist()] if '达标率' in df.columns else []
        rpm = [float(v) for v in df['转速(RPM)'].tolist()] if '转速(RPM)' in df.columns else []

        vehicle_types = df['车型'].unique().tolist() if '车型' in df.columns else []
        current_vehicle = vehicle_types[0] if vehicle_types else "未知车型"

        # ---------- KPI 计算 ----------
        avg_airflow = round(sum(airflow) / len(airflow), 2) if airflow else 0
        avg_noise = round(sum(noise) / len(noise), 2) if noise else 0
        avg_pass_rate = round(sum(pass_rate) / len(pass_rate), 2) if pass_rate else 0
        pass_count = sum(1 for p in pass_rate if p >= 60) if pass_rate else 0
        pass_ratio = round(pass_count / len(pass_rate) * 100, 1) if pass_rate else 0
        total_samples = len(fan_labels)

        # ---------- 异常检测 ----------
        noise_anomaly_z = detect_anomalies_zscore(noise, zscore_threshold)
        airflow_anomaly_z = detect_anomalies_zscore(airflow, zscore_threshold)
        noise_anomaly_iqr = detect_anomalies_iqr(noise)

        # 综合异常标记（Z-Score 或 IQR 任一命中即标记）
        anomaly_flags = []
        for i in range(total_samples):
            flags = []
            if noise_anomaly_z[i]:
                flags.append('噪音偏高(Z)')
            if airflow_anomaly_z[i]:
                flags.append('风量异常(Z)')
            if noise_anomaly_iqr[i]:
                flags.append('噪音异常(IQR)')
            if noise_alert_threshold > 0 and noise[i] > noise_alert_threshold:
                flags.append(f'噪音超标(>{noise_alert_threshold}dB)')
            if airflow_alert_threshold > 0 and airflow[i] < airflow_alert_threshold:
                flags.append(f'风量不足(<{airflow_alert_threshold}CFM)')
            anomaly_flags.append(flags)

        anomaly_count = sum(1 for f in anomaly_flags if f)
        anomaly_indices = [i for i, f in enumerate(anomaly_flags) if f]
        anomaly_details = [
            {
                'index': i,
                'label': fan_labels[i],
                'airflow': airflow[i],
                'noise': noise[i],
                'pass_rate': pass_rate[i] if i < len(pass_rate) else 0,
                'flags': anomaly_flags[i]
            }
            for i in anomaly_indices
        ]

        # ---------- 线性模型（供 What-if 使用） ----------
        linear_model = None
        if rpm and airflow and noise and '转速(RPM)' in df.columns:
            try:
                linear_model = fit_linear_model(df)
            except Exception:
                linear_model = None

        # ---------- KPI 卡片数据 ----------
        kpi = {
            'total_samples': total_samples,
            'avg_airflow': avg_airflow,
            'avg_noise': avg_noise,
            'avg_pass_rate': avg_pass_rate,
            'pass_ratio': pass_ratio,
            'anomaly_count': anomaly_count,
            'max_noise': round(max(noise), 1) if noise else 0,
            'min_noise': round(min(noise), 1) if noise else 0,
            'max_airflow': round(max(airflow), 1) if airflow else 0,
            'min_airflow': round(min(airflow), 1) if airflow else 0,
        }

        # ---------- AI 分析 ----------
        ai_analysis = "尚未接入 AI 大模型。请在页面顶端输入阿里云通义千问 API Key 进行智能分析。"

        if api_key and airflow and noise:
            try:
                noise_data = list(zip(fan_labels, noise))
                noise_data.sort(key=lambda x: x[1], reverse=True)
                worst_fans = [f"{item[0]}({item[1]}dB)" for item in noise_data[:3]]

                anomaly_summary = ""
                if anomaly_count > 0:
                    anomaly_summary = f"\n- ⚠️ 异常样本数: {anomaly_count}/{total_samples}"
                    anomaly_summary += f"\n- 异常详情: {', '.join([d['label'] + '(' + ','.join(d['flags']) + ')' for d in anomaly_details[:5]])}"

                prompt = f"""
你是一名资深的汽车NVH与气动流体力学专家。
当前你正在评估一张关于【{current_vehicle}】的扇叶风洞测试图表。
图表数据特征提取如下：
- 测试样本总数: {total_samples} 个风扇型号
- 整体平均风量: {avg_airflow} CFM
- 整体平均噪音: {avg_noise} dB
- 达标率: {pass_ratio}%（≥60%为合格）
- 噪音最大/最小值: {kpi['max_noise']}/{kpi['min_noise']} dB
- 风量最大/最小值: {kpi['max_airflow']}/{kpi['min_airflow']} CFM
- 图表中噪音表现最差（波峰）的三款型号: {', '.join(worst_fans)}
{anomaly_summary}

请完成以下输出：
1. **图表数据分析**：一句话总结当前批次风扇的风量与噪音呈现的整体趋势或暴露出的异常。
2. **异常诊断**：如果存在异常样本，分析可能的根因（设计缺陷？来料不良？工况偏离？）。
3. **车型特性匹配**：简要说明【{current_vehicle}】在实际工程应用中对散热和 NVH（噪音/振动）的核心诉求。
4. **专属优化建议**：针对上述高噪音的极值扇叶，给出 3 点具体的物理与空气动力学优化方案（例如：叶尖间隙、安装角、后缘锯齿等）。

要求：使用 Markdown 格式排版，条理清晰，字数控制在 500 字左右，直接输出专业干货。
"""
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}"
                }
                payload = {
                    "model": api_model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.7
                }
                response = requests.post(api_url, json=payload, headers=headers, timeout=300)
                if response.status_code == 200:
                    ai_analysis = response.json()['choices'][0]['message']['content']
                else:
                    ai_analysis = f"**AI 请求失败:** HTTP {response.status_code}。请检查 API Key 是否有效。"
            except Exception as e:
                ai_analysis = f"**调用异常:** {str(e)}"

        return jsonify({
            'success': True,
            'data': {
                'fan_labels': fan_labels,
                'airflow': airflow,
                'noise': noise,
                'pass_rate': pass_rate,
                'rpm': rpm,
                'anomaly_flags': anomaly_flags,
                'anomaly_details': anomaly_details,
                'anomaly_count': anomaly_count,
                'kpi': kpi,
                'linear_model': linear_model,
                'ai_analysis': ai_analysis,
                'current_vehicle': current_vehicle,
                'vehicle_types': vehicle_types,
            }
        })

    except Exception as e:
        return jsonify({'error': f'数据处理发生异常: {str(e)}'})


@app.route('/api/report', methods=['POST'])
def generate_report():
    """生成 Word 报告并返回下载"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': '缺少报告数据'})

        kpi = data.get('kpi', {})
        ai_analysis = data.get('ai_analysis', '')
        anomaly_details = data.get('anomaly_details', [])
        chart_image_b64 = data.get('chart_image', '')  # base64 ECharts 截图
        current_vehicle = data.get('current_vehicle', '未知车型')
        report_title = data.get('report_title', f'{current_vehicle} 风扇 NVH 测试报告')

        doc = Document()

        # ── 页面设置 ──
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        # ── 封面标题 ──
        doc.add_paragraph()  # 空行
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(report_title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('风扇 NVH 数据自动化评估报告')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f'车型: {current_vehicle}').font.size = Pt(11)
        info.add_run('\n')
        info.add_run(f'样本总数: {kpi.get("total_samples", "-")}').font.size = Pt(11)
        info.add_run('\n')
        info.add_run('生成日期: ' + pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')).font.size = Pt(11)

        doc.add_page_break()

        # ── 1. 概览摘要 ──
        doc.add_heading('一、测试概览', level=1)
        doc.add_paragraph(
            f'本次测试共覆盖 {kpi.get("total_samples", "-")} 个风扇型号。'
            f'整体平均风量为 {kpi.get("avg_airflow", "-")} CFM，'
            f'平均噪音为 {kpi.get("avg_noise", "-")} dB，'
            f'综合达标率为 {kpi.get("pass_ratio", "-")}%。'
        )

        # KPI 表格
        doc.add_heading('关键指标', level=2)
        table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        kpi_rows = [
            ('样本总数', str(kpi.get('total_samples', '-'))),
            ('平均风量', f'{kpi.get("avg_airflow", "-")} CFM'),
            ('平均噪音', f'{kpi.get("avg_noise", "-")} dB'),
            ('综合达标率', f'{kpi.get("pass_ratio", "-")}%'),
            ('异常样本数', str(kpi.get('anomaly_count', '-'))),
            ('噪音范围', f'{kpi.get("min_noise", "-")} ~ {kpi.get("max_noise", "-")} dB'),
        ]
        for i, (label, value) in enumerate(kpi_rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value

        # ── 2. 图表 ──
        if chart_image_b64:
            doc.add_heading('二、数据可视化图表', level=1)
            # 解码 base64 并保存为临时图片
            img_data = base64.b64decode(chart_image_b64.split(',')[1] if ',' in chart_image_b64 else chart_image_b64)
            img_path = os.path.join(app.config['UPLOAD_FOLDER'], '_report_chart.png')
            with open(img_path, 'wb') as f:
                f.write(img_data)
            doc.add_picture(img_path, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            os.remove(img_path)

        # ── 3. 异常详情 ──
        if anomaly_details:
            doc.add_heading('三、异常样本清单', level=1)
            anomaly_table = doc.add_table(rows=len(anomaly_details) + 1, cols=4, style='Light Grid Accent 1')
            anomaly_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = ['风扇型号', '风量(CFM)', '噪音(dB)', '异常标记']
            for j, h in enumerate(headers):
                anomaly_table.cell(0, j).text = h
                anomaly_table.cell(0, j).paragraphs[0].runs[0].bold = True
            for i, ad in enumerate(anomaly_details):
                anomaly_table.cell(i + 1, 0).text = ad.get('label', '')
                anomaly_table.cell(i + 1, 1).text = str(ad.get('airflow', ''))
                anomaly_table.cell(i + 1, 2).text = str(ad.get('noise', ''))
                anomaly_table.cell(i + 1, 3).text = ', '.join(ad.get('flags', []))

        # ── 4. AI 分析 ──
        if ai_analysis and '尚未接入' not in ai_analysis and '请求失败' not in ai_analysis and '调用异常' not in ai_analysis:
            doc.add_heading('四、AI 专家分析', level=1)
            # 简单处理 Markdown → 纯文本段落
            for line in ai_analysis.split('\n'):
                stripped = line.strip()
                if stripped.startswith('## '):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith('# '):
                    doc.add_heading(stripped[2:], level=2)
                elif stripped.startswith('- '):
                    doc.add_paragraph(stripped[2:], style='List Bullet')
                elif stripped.startswith('**') and '**' in stripped[4:]:
                    doc.add_paragraph(stripped.replace('**', ''))
                elif stripped:
                    doc.add_paragraph(stripped)

        # ── 5. 页脚 ──
        doc.add_paragraph()
        doc.add_paragraph('— 本报告由扇叶 NVH 数据自动化评估智能体自动生成 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe_filename = f"NVH_Report_{current_vehicle}_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}.docx"
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=safe_filename
        )

    except Exception as e:
        return jsonify({'error': f'报告生成失败: {str(e)}'})


@app.route('/api/optimize', methods=['POST'])
def optimize():
    """多目标优化接口：帕累托前沿 + AI 参数优化建议"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': '缺少数据'})

        fan_labels = data.get('fan_labels', [])
        airflow = data.get('airflow', [])
        noise = data.get('noise', [])
        rpm = data.get('rpm', [])
        pass_rate = data.get('pass_rate', [])
        current_vehicle = data.get('current_vehicle', '未知车型')
        noise_constraint = float(data.get('noise_constraint', 0))  # 噪音上限
        airflow_constraint = float(data.get('airflow_constraint', 0))  # 风量下限

        api_key = data.get('api_key', '').strip()
        api_model = data.get('api_model', 'qwen3.7-plus').strip()
        api_url = data.get('api_url',
            'https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions').strip()

        n = len(fan_labels)
        if n == 0:
            return jsonify({'error': '数据为空'})

        # ---------- 帕累托前沿 ----------
        pareto_flags = compute_pareto_front(airflow, noise)
        pareto_points = [
            {
                'label': fan_labels[i],
                'airflow': airflow[i],
                'noise': noise[i],
                'rpm': rpm[i] if i < len(rpm) else 0,
                'pass_rate': pass_rate[i] if i < len(pass_rate) else 0,
            }
            for i in range(n) if pareto_flags[i]
        ]
        # 按风量升序排列，确保前端连线正确
        pareto_points.sort(key=lambda p: p['airflow'])

        # ---------- 约束筛选 ----------
        feasible = []
        for i in range(n):
            ok = True
            if noise_constraint > 0 and noise[i] > noise_constraint:
                ok = False
            if airflow_constraint > 0 and airflow[i] < airflow_constraint:
                ok = False
            if ok:
                feasible.append({
                    'label': fan_labels[i],
                    'airflow': airflow[i],
                    'noise': noise[i],
                    'rpm': rpm[i] if i < len(rpm) else 0,
                    'pass_rate': pass_rate[i] if i < len(pass_rate) else 0,
                })
        feasible.sort(key=lambda p: p['noise'])  # 按噪音升序

        # ---------- What-if 线性模型 ----------
        linear_model = None
        if rpm and airflow and noise:
            try:
                # 临时构建 DataFrame 用于拟合
                temp_df = pd.DataFrame({'转速(RPM)': rpm, '风量(CFM)': airflow, '噪音(dB)': noise})
                linear_model = fit_linear_model(temp_df)
            except Exception:
                linear_model = None

        # ---------- AI 优化建议 ----------
        ai_optimization = ""
        if api_key and pareto_points:
            try:
                pareto_summary = '\n'.join([
                    f"  - {p['label']}: 风量={p['airflow']}CFM, 噪音={p['noise']}dB"
                    for p in pareto_points[:8]
                ])
                constraint_desc = ""
                if noise_constraint > 0:
                    constraint_desc += f"\n- 噪音上限约束: ≤{noise_constraint} dB"
                if airflow_constraint > 0:
                    constraint_desc += f"\n- 风量下限约束: ≥{airflow_constraint} CFM"
                feasible_count = len(feasible)
                constraint_desc += f"\n- 满足约束的可行方案数: {feasible_count}"

                prompt = f"""
你是一名资深的汽车散热风扇 NVH 优化专家。
当前车型: {current_vehicle}

帕累托前沿（风量-噪音最优权衡曲线）上的风扇型号：
{pareto_summary}

约束条件：{constraint_desc}

请完成以下输出：
1. **帕累托前沿解读**：分析帕累托前沿的形状特点，说明在哪个风量区间噪音上升最快（边际代价）。
2. **最优方案推荐**：从帕累托前沿中推荐 2-3 个最佳平衡点，并说明理由。
3. **参数优化路径**：基于当前最优方案，给出 3 条具体的物理参数调整建议以达到更好的风量-噪音平衡。包括但不限于：
   - 叶尖间隙调整（建议数值范围）
   - 叶片安装角优化（建议角度范围）
   - 后缘锯齿/波浪设计（建议几何参数）
   - 叶片数调整建议
   - 轮毂比优化

要求：Markdown 格式，条理清晰，建议具体可量化，字数控制在 450 字左右。
"""
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}"
                }
                payload = {
                    "model": api_model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.7
                }
                response = requests.post(api_url, json=payload, headers=headers, timeout=300)
                if response.status_code == 200:
                    ai_optimization = response.json()['choices'][0]['message']['content']
                else:
                    ai_optimization = f"**AI 请求失败:** HTTP {response.status_code}"
            except Exception as e:
                ai_optimization = f"**调用异常:** {str(e)}"

        return jsonify({
            'success': True,
            'data': {
                'pareto_points': pareto_points,
                'pareto_count': len(pareto_points),
                'feasible_solutions': feasible,
                'feasible_count': len(feasible),
                'linear_model': linear_model,
                'ai_optimization': ai_optimization,
            }
        })

    except Exception as e:
        return jsonify({'error': f'优化分析异常: {str(e)}'})


if __name__ == '__main__':
    app.run(debug=True, port=5000)
